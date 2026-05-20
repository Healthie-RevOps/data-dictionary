#!/usr/bin/env python3
"""Generate Healthie HubSpot Data Dictionary by querying live HubSpot CRM metadata.

Produces Healthie_HubSpot_Data_Dictionary_v1.docx mirroring the structure of
Healthie_SFDC_Data_Dictionary_v3.docx (used as a style template).

Stdlib-only (no `requests`, no `PyYAML`). Uses python-docx for the .docx output.
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_DIR = Path(__file__).resolve().parent
CACHE_DIR = PROJECT_DIR / ".dictionary_cache"
HS_CONFIG_PATH = PROJECT_DIR / "hubspot.config.yml"
SFDC_TEMPLATE = PROJECT_DIR / "Healthie_SFDC_Data_Dictionary_v3.docx"
OUTPUT_DOCX = PROJECT_DIR / "Healthie_HubSpot_Data_Dictionary_v2.docx"
PORTAL_NAME = "healthie-prod"
PORTAL_ID = 43826161
API_BASE = "https://api.hubapi.com"
THROTTLE_RPS = 9
CALL_DELAY_S = 1.0 / THROTTLE_RPS

OBJECT_TYPES = [
    "contacts", "companies", "deals", "tickets", "feedback_submissions", "leads",
    "products", "line_items", "quotes", "invoices", "subscriptions",
    "payments", "orders", "carts", "discounts", "fees", "taxes", "commerce_payments",
    "calls", "emails", "meetings", "notes", "tasks", "communications", "postal_mail",
    "marketing_events",
]

OBJECT_LABEL_OVERRIDES = {
    "contacts": "Contacts",
    "companies": "Companies",
    "deals": "Deals",
    "tickets": "Tickets",
    "feedback_submissions": "Feedback Submissions",
    "leads": "Leads",
    "products": "Products",
    "line_items": "Line Items",
    "quotes": "Quotes",
    "invoices": "Invoices",
    "subscriptions": "Subscriptions",
    "payments": "Payments",
    "orders": "Orders",
    "carts": "Carts",
    "discounts": "Discounts",
    "fees": "Fees",
    "taxes": "Taxes",
    "commerce_payments": "Commerce Payments",
    "calls": "Calls",
    "emails": "Emails",
    "meetings": "Meetings",
    "notes": "Notes",
    "tasks": "Tasks",
    "communications": "Communications",
    "postal_mail": "Postal Mail",
    "marketing_events": "Marketing Events",
}

OBJECT_ORDER = [
    "contacts", "companies", "deals", "tickets", "leads", "feedback_submissions",
    "products", "line_items", "quotes", "invoices", "subscriptions",
    "payments", "orders", "carts", "discounts", "fees", "taxes", "commerce_payments",
    "marketing_events",
    "calls", "emails", "meetings", "notes", "tasks", "communications", "postal_mail",
]


def read_access_token() -> str:
    text = HS_CONFIG_PATH.read_text()
    m = re.search(r"accessToken:\s*>-\s*\n\s*(\S+)", text)
    if not m:
        sys.exit("Could not find accessToken in hubspot.config.yml")
    return m.group(1)


_last_call_ts = 0.0


def _throttle() -> None:
    global _last_call_ts
    delta = time.time() - _last_call_ts
    if delta < CALL_DELAY_S:
        time.sleep(CALL_DELAY_S - delta)
    _last_call_ts = time.time()


def hs_request(method: str, path: str, token_box: list[str], body=None, retries: int = 3):
    """HTTP request with throttle, 429 backoff, and 401-token-refresh.

    `token_box` is a single-element list so a refreshed token is shared across calls.
    Returns (status_code, parsed_json_or_dict).
    """
    url = f"{API_BASE}{path}"
    data = json.dumps(body).encode() if body is not None else None
    for attempt in range(retries):
        _throttle()
        token = token_box[0]
        headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
        req = urllib.request.Request(url, data=data, headers=headers, method=method)
        try:
            with urllib.request.urlopen(req, timeout=60) as resp:
                return resp.status, json.loads(resp.read().decode())
        except urllib.error.HTTPError as e:
            try:
                err_body = json.loads(e.read().decode())
            except Exception:
                err_body = {"error": str(e)}
            if e.code == 429:
                wait = int(e.headers.get("Retry-After", "5"))
                time.sleep(wait)
                continue
            if e.code == 401 and attempt < retries - 1:
                fresh = read_access_token()
                if fresh != token:
                    token_box[0] = fresh
                    continue
                return e.code, err_body
            return e.code, err_body
        except Exception as e:
            if attempt == retries - 1:
                return 0, {"error": str(e)}
            time.sleep(2 ** attempt)
    return 0, {"error": "exhausted retries"}


def cache_path(name: str) -> Path:
    CACHE_DIR.mkdir(exist_ok=True)
    return CACHE_DIR / name


def cache_load(name: str):
    p = cache_path(name)
    if p.exists():
        try:
            return json.loads(p.read_text())
        except Exception:
            return None
    return None


def cache_save(name: str, value) -> None:
    cache_path(name).write_text(json.dumps(value, indent=2))


def discover_objects(token_box: list[str], refresh: bool = False) -> list[dict]:
    cache_name = "objects_discovered.json"
    if not refresh:
        cached = cache_load(cache_name)
        if cached is not None:
            print(f"[discover] using cached probes ({len(cached)} entries)")
            return cached
    results: list[dict] = []
    for obj in OBJECT_TYPES:
        status, data = hs_request("GET", f"/crm/v3/properties/{obj}", token_box)
        prop_count = len(data.get("results", [])) if status == 200 else 0
        entry = {"name": obj, "label": OBJECT_LABEL_OVERRIDES.get(obj, obj), "status": status, "property_count": prop_count}
        if status != 200:
            entry["error"] = (data.get("message") or data.get("category") or str(data))[:200]
        results.append(entry)
        print(f"[discover] {obj}: HTTP {status} ({prop_count} props)")
    # custom objects via schemas (best effort)
    status, data = hs_request("GET", "/crm/v3/schemas", token_box)
    if status == 200:
        for schema in data.get("results", []):
            otid = schema.get("objectTypeId") or ""
            if not otid.startswith("2-"):
                continue
            pname = schema.get("name") or otid
            label = (schema.get("labels") or {}).get("plural") or pname
            pstatus, pdata = hs_request("GET", f"/crm/v3/properties/{otid}", token_box)
            entry = {
                "name": otid,
                "internal_name": pname,
                "label": label,
                "status": pstatus,
                "property_count": len(pdata.get("results", [])) if pstatus == 200 else 0,
                "is_custom": True,
            }
            results.append(entry)
            print(f"[discover] custom {otid} ({pname}): HTTP {pstatus} ({entry['property_count']} props)")
    else:
        msg = (data.get("message") or str(data))[:200]
        print(f"[discover] /crm/v3/schemas returned {status}; custom-object discovery unavailable. ({msg})")
        results.append({
            "name": "_schemas_endpoint",
            "status": status,
            "note": "Custom object discovery requires crm.schemas.custom.read scope.",
            "error": msg,
        })
    cache_save(cache_name, results)
    return results


def fetch_properties(obj_name: str, token_box: list[str], refresh: bool = False) -> list[dict]:
    cache_name = f"props_{obj_name}.json"
    if not refresh:
        cached = cache_load(cache_name)
        if cached is not None:
            return cached
    status, data = hs_request("GET", f"/crm/v3/properties/{obj_name}", token_box)
    if status != 200:
        return []
    props = data.get("results", [])
    cache_save(cache_name, props)
    return props


def fetch_total_count(obj_name: str, token_box: list[str], refresh: bool = False) -> int | None:
    cache_name = f"count_{obj_name}.json"
    if not refresh:
        cached = cache_load(cache_name)
        if cached is not None:
            return cached.get("total")
    status, data = hs_request(
        "POST", f"/crm/v3/objects/{obj_name}/search", token_box,
        body={"limit": 1, "properties": []},
    )
    if status != 200:
        cache_save(cache_name, {"total": None, "status": status, "error": data})
        return None
    total = data.get("total")
    cache_save(cache_name, {"total": total})
    return total


def fetch_pop_counts(obj_name: str, prop_names: list[str], token_box: list[str], refresh: bool = False) -> dict[str, int | None]:
    cache_name = f"pop_{obj_name}.json"
    cached = {} if refresh else (cache_load(cache_name) or {})
    out = dict(cached)
    todo = [p for p in prop_names if p not in out]
    if not todo:
        return out
    print(f"[pop] {obj_name}: {len(todo)} property counts to fetch (cached: {len(cached)})")
    for i, prop in enumerate(todo, 1):
        status, data = hs_request(
            "POST", f"/crm/v3/objects/{obj_name}/search", token_box,
            body={
                "filterGroups": [{"filters": [{"propertyName": prop, "operator": "HAS_PROPERTY"}]}],
                "limit": 1,
                "properties": [],
            },
        )
        if status == 200:
            out[prop] = data.get("total")
        else:
            out[prop] = None
        if i % 25 == 0:
            cache_save(cache_name, out)
            print(f"[pop] {obj_name}: {i}/{len(todo)}")
    cache_save(cache_name, out)
    return out


# ---------- Heuristics ----------

PII_PAT = re.compile(
    r"(^|_)(email|phone|mobile|fax|address|firstname|first_name|lastname|last_name|fullname|full_name|"
    r"zip|postal|city|country|state|ip_address|dob|birth|ssn|hs_calculated_phone|hs_searchable_calculated_phone)(_|$)",
    re.I,
)
FINANCIAL_ID_PAT = re.compile(r"(stripe|payment_id|invoice_id|subscription_id|charge_id|hs_object_id)", re.I)
FINANCIAL_PAT = re.compile(
    r"(^|_)(amount|revenue|arr|mrr|tcv|acv|deal_value|price|total|fee|salary|budget|payment|cost|spend)(_|$)",
    re.I,
)


def classify_sensitivity(prop: dict) -> str:
    name = prop.get("name", "")
    label = prop.get("label", "")
    if name == "hs_object_id":
        return "Internal"
    if PII_PAT.search(name) or PII_PAT.search(label):
        return "PII"
    if FINANCIAL_ID_PAT.search(name):
        return "Financial ID"
    if FINANCIAL_PAT.search(name) or FINANCIAL_PAT.search(label):
        return "Financial"
    return "Internal"


MARKETING_GROUP_PAT = re.compile(
    r"(analytics|conversion|^email$|^emailinformation$|socialmedia|_ads_|formfield|calendly|campaign|seventh_sense|chameleon|"
    r"intent|webinar|event_attendance|gotowebinar)",
    re.I,
)
SALES_GROUP_PAT = re.compile(r"(sales_propert|deal_|pipeline|opportunity|^dealinfo)", re.I)
SUPPORT_GROUP_PAT = re.compile(r"(service_|support_|ticket_|conversation_|knowledge_)", re.I)


def classify_owner(prop: dict, obj_name: str, sensitivity: str) -> str:
    group = (prop.get("groupName") or "").lower()
    if obj_name == "deals":
        return "Sales"
    if obj_name == "tickets":
        return "Customer Success"
    if obj_name == "marketing_events":
        return "Marketing"
    if obj_name in {"calls", "emails", "meetings", "notes", "tasks", "communications", "postal_mail"}:
        return "Sales / RevOps"
    if MARKETING_GROUP_PAT.search(group):
        return "Marketing"
    if SALES_GROUP_PAT.search(group):
        return "Sales"
    if SUPPORT_GROUP_PAT.search(group):
        return "Customer Success"
    if sensitivity in {"Financial", "Financial ID"}:
        return "Finance"
    return "RevOps"


def classify_source(prop: dict) -> str:
    name = prop.get("name", "")
    group = (prop.get("groupName") or "").lower()
    calculated = bool(prop.get("calculated"))
    hs_defined = bool(prop.get("hubspotDefined"))
    form_field = bool(prop.get("formField"))
    if calculated:
        return "HubSpot calculated"
    if group == "salesforceinformation":
        return "Salesforce sync"
    if name.startswith("hs_") or hs_defined:
        return "HubSpot system"
    if form_field:
        return "Form submission"
    return "Manual entry (or integration)"


TYPE_LABEL_BY_FIELD = {
    "text": "Single-line text",
    "textarea": "Multi-line text",
    "number": "Number",
    "select": "Dropdown",
    "radio": "Radio select",
    "checkbox": "Checkbox",
    "booleancheckbox": "Boolean checkbox",
    "date": "Date",
    "file": "File",
    "phonenumber": "Phone",
    "html": "Rich text",
    "calculation_equation": "Calculation",
    "calculation_score": "Score",
    "calculation_rollup": "Rollup",
    "calculation_read_time": "Read time",
}
TYPE_LABEL_BY_TYPE = {
    "string": "Text",
    "number": "Number",
    "bool": "Boolean",
    "datetime": "Date/Time",
    "date": "Date",
    "enumeration": "Enumeration",
    "phone_number": "Phone",
    "object_coordinates": "Object reference",
    "json": "JSON",
}


def humanize_type(prop: dict) -> str:
    field_type = prop.get("fieldType") or ""
    base_type = prop.get("type") or ""
    if field_type in TYPE_LABEL_BY_FIELD:
        label = TYPE_LABEL_BY_FIELD[field_type]
    elif base_type in TYPE_LABEL_BY_TYPE:
        label = TYPE_LABEL_BY_TYPE[base_type]
    else:
        label = (base_type or field_type or "Unknown").replace("_", " ").title()
    if prop.get("calculated"):
        label = f"Calculated ({label})"
    if prop.get("referencedObjectType"):
        label = f"Reference → {prop['referencedObjectType']}"
    return label


def derive_description(prop: dict, source: str) -> str:
    desc = (prop.get("description") or "").strip()
    if desc:
        return desc
    label = prop.get("label") or prop.get("name") or "field"
    type_str = humanize_type(prop)
    if prop.get("type") == "enumeration":
        opts = [o.get("label") or o.get("value") for o in (prop.get("options") or []) if not o.get("hidden")]
        opts = [o for o in opts if o]
        if opts:
            preview = ", ".join(opts[:6])
            more = "" if len(opts) <= 6 else f", … ({len(opts) - 6} more)"
            return f"{type_str} — {label}. (Source: {source}.) Active values: {preview}{more}."
    return f"{type_str} — {label}. (Source: {source}.)"


def derive_dependencies(prop: dict, obj_name: str) -> str:
    notes: list[str] = []
    if prop.get("calculated"):
        notes.append("Calculated by HubSpot — see portal for formula.")
    if (prop.get("groupName") or "").lower() == "salesforceinformation":
        notes.append("Synced with Salesforce; see SFDC dictionary for paired field.")
    if prop.get("referencedObjectType"):
        notes.append(f"References {prop['referencedObjectType']}.")
    if prop.get("hasUniqueValue"):
        notes.append("Unique values enforced.")
    return " ".join(notes) if notes else "—"


GROUP_LABEL_OVERRIDES = {
    "contactinformation": "Contact Information",
    "companyinformation": "Company Information",
    "dealinformation": "Deal Information",
    "salesforceinformation": "Salesforce Information",
    "emailinformation": "Email Information",
    "conversioninformation": "Conversion Information",
    "analyticsinformation": "Analytics Information",
    "socialmediainformation": "Social Media Information",
    "contact_activity": "Contact Activity",
    "contactlcs": "Lifecycle / Status",
    "companylcs": "Company Lifecycle / Status",
    "deal_information": "Deal Information",
    "dealinformation_lcs": "Deal Lifecycle / Status",
    "sales_properties": "Sales Properties",
    "facebook_ads_properties": "Facebook Ads",
    "lead_ads": "Lead Ads",
    "google_ads": "Google Ads",
    "linkedin_ads": "LinkedIn Ads",
    "calendly": "Calendly",
    "zoom": "Zoom",
    "seventh_sense": "Seventh Sense",
    "intentinformation": "Intent Data",
    "engagement": "Engagement",
    "core": "Core",
    "other": "Other",
    "membershipinformation": "Membership",
    "order_information": "Order Information",
    "deal_revenue": "Deal Revenue",
    "post_sale": "Post Sale",
    "subscription": "Subscription",
    "engagementinformation": "Engagement Information",
    "videos": "Videos",
}


def humanize_group(group_name: str) -> str:
    if not group_name:
        return "Other"
    key = group_name.lower()
    if key in GROUP_LABEL_OVERRIDES:
        return GROUP_LABEL_OVERRIDES[key]
    base = re.sub(r"[_\-]+", " ", group_name).strip()
    base = re.sub(r"(?<!^)(?=[A-Z])", " ", base)
    base = re.sub(r"(?i)(\w+?)(information|properties|details|info)$", r"\1 \2", base)
    base = re.sub(r"\s+", " ", base).strip()
    return " ".join(w[:1].upper() + w[1:] for w in base.split())


def format_pop(count: int | None, total: int | None) -> str:
    if total is None or total == 0 or count is None:
        return "—"
    pct = round(100.0 * count / total)
    pct = max(0, min(100, pct))
    suffix = "*" if total < 100 else ""
    return f"{pct}%{suffix}"


# ---------- Build dictionary model ----------

def build_object_section(obj: dict, props: list[dict], pop: dict, total: int | None) -> dict:
    rows_by_group: dict[str, list[dict]] = {}
    enumeration_appendix: list[tuple[str, str, list[dict]]] = []  # (object_label, prop_label, options)
    for p in props:
        if p.get("archived") or p.get("hidden"):
            continue
        sensitivity = classify_sensitivity(p)
        owner = classify_owner(p, obj["name"], sensitivity)
        source = classify_source(p)
        description = derive_description(p, source)
        dependencies = derive_dependencies(p, obj["name"])
        pop_str = format_pop(pop.get(p["name"]), total)
        group = p.get("groupName") or "other"
        rows_by_group.setdefault(group, []).append({
            "label": p.get("label") or p.get("name"),
            "name": p.get("name"),
            "type": humanize_type(p),
            "pop": pop_str,
            "source": source,
            "owner": owner,
            "sensitivity": sensitivity,
            "description": description,
            "dependencies": dependencies,
        })
        if p.get("type") == "enumeration":
            opts = [o for o in (p.get("options") or []) if not o.get("hidden")]
            if opts:
                enumeration_appendix.append((obj["label"], p.get("label") or p.get("name"), opts))

    ordered_groups = sorted(rows_by_group.items(), key=lambda kv: (-len(kv[1]), kv[0]))
    sections = [{"group": humanize_group(g), "raw_group": g, "rows": rows} for g, rows in ordered_groups]
    return {
        "name": obj["name"],
        "label": obj["label"],
        "total_records": total,
        "property_count_visible": sum(len(s["rows"]) for s in sections),
        "property_count_total": len(props),
        "sections": sections,
        "enumerations": enumeration_appendix,
    }


# ---------- DOCX rendering ----------

FIELD_TABLE_HEADERS = [
    "Field Label", "Data Type", "Pop %", "Source / Update",
    "Owner", "Sensitivity", "Description", "Dependencies / Notes",
]


def _clear_body(doc: Document) -> None:
    body = doc.element.body
    sectPr = body.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"
    )
    for child in list(body):
        if child is sectPr:
            continue
        body.remove(child)


def _add_para(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
    return p


def _try_style(doc: Document, candidates: list[str]) -> str | None:
    available = {s.name for s in doc.styles}
    for c in candidates:
        if c in available:
            return c
    return None


# ---- V3 visual style constants ----
HEADER_FILL_HEX = "1f4e79"
HEADER_TEXT_HEX = "ffffff"
TABLE_BORDER_HEX = "000000"
CELL_BORDER_HEX = "bfbfbf"
HEADER_FONT_HALFPT = 18  # 9pt
BODY_FONT_HALFPT = 18    # 9pt body for compact tables; V3 uses similar

# Column-width recipes (dxa) keyed by (section_kind, col_count)
COL_WIDTHS = {
    ("landscape", 8): [1700, 1300, 700, 1900, 1300, 1100, 3380, 2300],   # field tables
    ("landscape", 4): [3000, 1500, 1500, 7680],                           # discovery results
    ("landscape", 3): [1700, 6280, 5700],                                 # ad hoc 3-col landscape
    ("landscape", 2): [3500, 10180],
    ("portrait", 2):  [2400, 6960],
    ("portrait", 3):  [2200, 4400, 2760],
    ("portrait", 4):  [1700, 1900, 2000, 3760],
}
TOTAL_WIDTH = {"landscape": 13680, "portrait": 9360}


def _orient_label(doc: Document) -> str:
    section = doc.sections[-1]
    if section.orientation == WD_ORIENT.LANDSCAPE:
        return "landscape"
    return "portrait"


def _set_tbl_borders(table) -> None:
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._element.insert(0, tblPr)
    for prev in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(prev)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), TABLE_BORDER_HEX)
        borders.append(e)
    tblPr.append(borders)


def _set_tbl_layout_fixed(table, total_width_dxa: int) -> None:
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._element.insert(0, tblPr)
    for prev in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(prev)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for prev in tblPr.findall(qn("w:tblW")):
        tblPr.remove(prev)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(total_width_dxa))
    tblPr.append(tblW)
    for prev in tblPr.findall(qn("w:tblLook")):
        tblPr.remove(prev)
    tblLook = OxmlElement("w:tblLook")
    tblLook.set(qn("w:val"), "0000")
    tblPr.append(tblLook)


def _set_grid_cols(table, widths_dxa: list[int]) -> None:
    grid = table._element.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        # insert grid after tblPr
        tblPr = table._element.find(qn("w:tblPr"))
        if tblPr is not None:
            tblPr.addnext(grid)
        else:
            table._element.insert(0, grid)
    for prev in list(grid):
        grid.remove(prev)
    for w in widths_dxa:
        c = OxmlElement("w:gridCol")
        c.set(qn("w:w"), str(w))
        grid.append(c)


def _set_cell_width(cell, width_dxa: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for prev in tcPr.findall(qn("w:tcW")):
        tcPr.remove(prev)
    w = OxmlElement("w:tcW")
    w.set(qn("w:type"), "dxa")
    w.set(qn("w:w"), str(width_dxa))
    tcPr.append(w)


def _set_cell_borders(cell, color_hex: str = CELL_BORDER_HEX) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for prev in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(prev)
    bd = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color_hex)
        bd.append(e)
    tcPr.append(bd)


def _set_cell_margins(cell, top=80, left=100, bottom=80, right=100) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for prev in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(prev)
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)


def _set_cell_valign_top(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for prev in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(prev)
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), "top")
    tcPr.append(v)


def _shade_cell(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for prev in tcPr.findall(qn("w:shd")):
        tcPr.remove(prev)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _write_cell_text(cell, text: str, *, bold: bool = False, color_hex: str | None = None,
                      size_halfpt: int = BODY_FONT_HALFPT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text if text is not None else "")
    run.bold = bold
    run.font.size = Pt(size_halfpt / 2.0)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex.upper())
    # tighter paragraph spacing
    pPr = p._p.get_or_add_pPr()
    for prev in pPr.findall(qn("w:spacing")):
        pPr.remove(prev)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), "20")
    sp.set(qn("w:after"), "20")
    sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def _add_v3_table(doc: Document, headers: list[str], rows: list[list[str]],
                   *, col_widths: list[int] | None = None) -> None:
    n_cols = len(headers)
    orient = _orient_label(doc)
    if col_widths is None:
        col_widths = COL_WIDTHS.get((orient, n_cols))
    if col_widths is None:
        # equal split
        total = TOTAL_WIDTH[orient]
        each = total // n_cols
        col_widths = [each] * n_cols
    total_width = sum(col_widths)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    _set_tbl_borders(table)
    _set_tbl_layout_fixed(table, total_width)
    _set_grid_cols(table, col_widths)
    # header row
    for c_i, h in enumerate(headers):
        cell = table.rows[0].cells[c_i]
        _set_cell_borders(cell)
        _set_cell_margins(cell)
        _set_cell_valign_top(cell)
        _set_cell_width(cell, col_widths[c_i])
        _shade_cell(cell, HEADER_FILL_HEX)
        _write_cell_text(cell, h, bold=True, color_hex=HEADER_TEXT_HEX, size_halfpt=HEADER_FONT_HALFPT)
    # data rows
    for r_i, row in enumerate(rows, start=1):
        for c_i in range(n_cols):
            val = row[c_i] if c_i < len(row) else ""
            cell = table.rows[r_i].cells[c_i]
            _set_cell_borders(cell)
            _set_cell_margins(cell)
            _set_cell_valign_top(cell)
            _set_cell_width(cell, col_widths[c_i])
            _write_cell_text(cell, val if val is not None else "", size_halfpt=BODY_FONT_HALFPT)


def _add_two_col_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Backwards-compatible name kept for existing callers; routes to V3 styling."""
    _add_v3_table(doc, headers, rows)


def _set_section_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)


def _set_section_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)


def render_docx(model: dict, run_meta: dict, out_path: Path) -> None:
    if SFDC_TEMPLATE.exists():
        doc = Document(str(SFDC_TEMPLATE))
        _clear_body(doc)
    else:
        doc = Document()

    # Initial section: portrait for front matter
    _set_section_portrait(doc.sections[-1])

    h1 = _try_style(doc, ["Heading 1"]) or "Heading 1"
    h2 = _try_style(doc, ["Heading 2"]) or "Heading 2"
    title = _try_style(doc, ["Title"]) or h1

    # Title
    _add_para(doc, "Healthie HubSpot Data Dictionary v1", style=title)
    _add_para(
        doc,
        f"Portal: {PORTAL_NAME} (id {PORTAL_ID}). Generated {run_meta['generated_at']}. "
        f"{run_meta['object_count']} objects, {run_meta['property_count']} live properties.",
    )

    # How to Use
    _add_para(doc, "How to Use This Dictionary", style=h1)

    _add_para(doc, "Methodology", style=h2)
    _add_para(
        doc,
        "This dictionary catalogs the live HubSpot properties on every object accessible "
        "from the production portal. For each property we record its label, HubSpot data "
        "type, current population, derived source-of-truth, suggested owner, sensitivity "
        "class, description, and dependencies. It mirrors the Healthie Salesforce Data "
        "Dictionary v3 structure so the two can be read in parallel.",
    )
    _add_para(
        doc,
        "Auto-derived columns (Source / Update, Owner, Sensitivity, Description, "
        "Dependencies / Notes) are heuristic starting points based on property metadata "
        "and naming conventions. Treat them as proposed fills for human review, not "
        "system-of-record statements.",
    )

    _add_para(doc, "Column Definitions", style=h2)
    _add_two_col_table(
        doc,
        ["Column", "What it captures"],
        [
            ["Field Label", "Display name in HubSpot."],
            ["Data Type", "HubSpot field type (and underlying type)."],
            ["Pop %", "Percent of records with this property populated. Computed as count(HAS_PROPERTY)/count(*) at run time. * suffix flags small populations (< 100 records)."],
            ["Source / Update", "How the value gets written: Manual entry, Form submission, Salesforce sync, HubSpot calculated, HubSpot system, or integration."],
            ["Owner", "Team that owns the field's quality and definition."],
            ["Sensitivity", "PII / Financial / Financial ID / Internal."],
            ["Description", "Property description from HubSpot when set; otherwise an auto-generated typed placeholder."],
            ["Dependencies / Notes", "Calculated formulas, Salesforce sync notes, references to other objects, uniqueness constraints, etc."],
        ],
    )

    _add_para(doc, "Review Cadence", style=h2)
    _add_para(
        doc,
        "Regenerate quarterly or after any major HubSpot integration or schema change. "
        "Re-run generate_hubspot_dictionary.py with --refresh to invalidate the metadata cache.",
    )

    # Governance
    _add_para(doc, "Governance", style=h1)
    _add_two_col_table(
        doc,
        ["Area", "Owner / Process"],
        [
            ["Dictionary maintenance", "RevOps (Director: Bill Coffin)."],
            ["Property creation in HubSpot", "RevOps approval; field group + description required."],
            ["Salesforce-synced fields", "Mirror SFDC dictionary; pair with QualifiedApiName entry there."],
            ["Sensitivity overrides", "Privacy/Security review for any PII/Financial reclassification."],
            ["Discrepancies vs. portal", "Reported via the Data Dictionary Slack channel; corrections trigger a new run."],
        ],
    )

    # Data Sources & Integrations
    _add_para(doc, "Data Sources & Integrations", style=h1)
    integration_rows = build_integration_rows(model)
    _add_two_col_table(
        doc, ["Source System", "Writes to", "Notes"], integration_rows,
    )

    # Glossary
    _add_para(doc, "Glossary", style=h1)
    _add_two_col_table(
        doc,
        ["Term", "Definition"],
        [
            ["Property", "A field on a HubSpot object. Equivalent to a Salesforce custom/standard field."],
            ["Property Group", "HubSpot's native grouping of related properties; drives the sub-section layout in this document."],
            ["Lifecycle Stage", "Stage of the contact/company in the marketing-to-sales journey (Subscriber → Lead → MQL → SQL → Customer)."],
            ["Lead Status", "Sub-stage within Lead/MQL/SQL; sales-managed."],
            ["Original Source", "First-touch attribution channel."],
            ["MQL", "Marketing Qualified Lead — meets marketing's bar for sales handoff."],
            ["SQL", "Sales Qualified Lead — accepted by sales after qualification."],
            ["Calculated property", "Server-side formula maintained by HubSpot. Read-only."],
            ["HubSpot-defined property", "Standard property shipped by HubSpot (vs. custom)."],
            ["Form field", "Property surfaced on at least one HubSpot form."],
            ["Pipeline / Stage", "Deal pipeline and the stages it contains; each deal lives in exactly one stage."],
            ["Ticket Pipeline / Status", "Equivalent for tickets."],
            ["Engagement", "Activity record (call, email, meeting, note, task) attached to a CRM object."],
            ["Custom object", "Portal-specific object with `2-xxxxx` objectTypeId."],
        ],
    )

    # Switch to landscape for the field-table-heavy object sections
    landscape_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_section_landscape(landscape_section)

    # Object sections
    enumerations_global: list[tuple[str, str, list[dict]]] = []
    for obj in model["objects"]:
        _add_para(doc, obj["label"], style=h1)
        _add_para(
            doc,
            f"{obj['property_count_visible']} live properties "
            f"({obj['property_count_total']} total including archived/hidden). "
            f"Total records: {obj['total_records'] if obj['total_records'] is not None else 'unknown'}.",
        )
        for sec in obj["sections"]:
            _add_para(doc, f"{sec['group']} ({len(sec['rows'])})", style=h2)
            rows = [
                [
                    r["label"], r["type"], r["pop"], r["source"],
                    r["owner"], r["sensitivity"], r["description"], r["dependencies"],
                ]
                for r in sec["rows"]
            ]
            _add_two_col_table(doc, FIELD_TABLE_HEADERS, rows)
        enumerations_global.extend(obj["enumerations"])

    # Switch back to portrait for the appendix and changelog
    appendix_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_section_portrait(appendix_section)

    # Appendix: Enumeration Values
    _add_para(doc, "Appendix: Enumeration Values", style=h1)
    _add_para(
        doc,
        "Active option values for every non-hidden enumeration property, grouped by object and field.",
    )
    by_object: dict[str, list[tuple[str, list[dict]]]] = {}
    for obj_label, prop_label, opts in enumerations_global:
        by_object.setdefault(obj_label, []).append((prop_label, opts))
    for obj_label in [o["label"] for o in model["objects"]]:
        entries = by_object.get(obj_label)
        if not entries:
            continue
        _add_para(doc, obj_label, style=h2)
        for prop_label, opts in entries:
            _add_para(doc, prop_label, style=_try_style(doc, ["Heading 3"]) or h2)
            _add_two_col_table(
                doc,
                ["Value", "Label", "Display Order"],
                [[o.get("value", ""), o.get("label", ""), str(o.get("displayOrder", ""))] for o in opts],
            )

    # Appendix: Discovered vs. Skipped Objects
    _add_para(doc, "Appendix: Object Discovery Results", style=h1)
    _add_two_col_table(
        doc,
        ["Object Type", "HTTP Status", "Property Count", "Notes"],
        [
            [
                d.get("name", ""),
                str(d.get("status", "")),
                str(d.get("property_count", "")),
                d.get("error") or d.get("note") or ("custom object" if d.get("is_custom") else "—"),
            ]
            for d in model["discovery"]
        ],
    )

    # Changelog
    _add_para(doc, "Changelog", style=h1)
    _add_two_col_table(
        doc,
        ["Version", "Date", "Notes"],
        [["v1", run_meta["generated_at_date"], "Initial generation. Auto-pulled from HubSpot CRM API; heuristic fills for Source/Update, Owner, Sensitivity, and Dependencies. Pop % computed live."]],
    )

    doc.save(str(out_path))


def build_integration_rows(model: dict) -> list[list[str]]:
    """Inspect property groups across objects to infer active integrations."""
    groups_seen: set[str] = set()
    for obj in model["objects"]:
        for sec in obj["sections"]:
            groups_seen.add(sec["raw_group"].lower())
    rows: list[list[str]] = []

    def has(*pats: str) -> bool:
        return any(any(re.search(pat, g) for pat in pats) for g in groups_seen)

    rows.append([
        "HubSpot system",
        "All objects (hs_* properties, calculated fields)",
        "HubSpot's built-in lifecycle, attribution, and analytics fields. Maintained by HubSpot, read-only from the user perspective.",
    ])
    if has(r"salesforceinformation"):
        rows.append([
            "Salesforce",
            "Contacts, Companies, Deals (salesforceinformation group)",
            "Bidirectional sync via HubSpot's Salesforce integration. See Healthie SFDC Data Dictionary v3 for the paired Salesforce field on each row.",
        ])
    if has(r"calendly"):
        rows.append(["Calendly", "Contacts (calendly group)", "Booking and meeting metadata."])
    if has(r"facebook_ads"):
        rows.append(["Facebook Ads", "Contacts (facebook_ads_properties)", "Lead form, ad, campaign IDs."])
    if has(r"google[_]?ads", r"adwords"):
        rows.append(["Google Ads", "Contacts (google_ads / adwords)", "Lead form, ad, campaign IDs."])
    if has(r"linkedin"):
        rows.append(["LinkedIn Ads", "Contacts (linkedin)", "Lead Gen Form metadata."])
    if has(r"zoom"):
        rows.append(["Zoom", "Contacts (zoom)", "Webinar and meeting attendance."])
    if has(r"gotowebinar", r"webinar"):
        rows.append(["Webinar platform", "Contacts (webinar groups)", "Registration and attendance."])
    if has(r"chameleon"):
        rows.append(["Chameleon", "Contacts (chameleon profile groups)", "Product-tour / in-app guidance signals."])
    if has(r"seventh_sense"):
        rows.append(["Seventh Sense", "Contacts (seventh_sense)", "Email send-time optimization."])
    if has(r"intent"):
        rows.append(["Intent data provider", "Contacts/Companies (intent groups)", "Third-party intent signals."])
    rows.append([
        "Forms (HubSpot)",
        "Contacts (formField properties)",
        "Any property surfaced on a HubSpot form. Captured on submission.",
    ])
    rows.append([
        "Manual entry",
        "All objects (default for non-system, non-form properties)",
        "Reps and ops users editing records directly in HubSpot.",
    ])
    return rows


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate Healthie HubSpot Data Dictionary v1.")
    parser.add_argument("--refresh", action="store_true", help="Invalidate all caches and re-fetch from HubSpot.")
    parser.add_argument("--refresh-discovery", action="store_true", help="Re-probe objects only.")
    parser.add_argument("--refresh-pop", action="store_true", help="Re-fetch population counts only.")
    parser.add_argument("--skip-pop", action="store_true", help="Skip Pop % queries; render with '—'.")
    parser.add_argument("--render-only", action="store_true", help="Skip all API calls; render from cache.")
    args = parser.parse_args()

    token = read_access_token()
    token_box = [token]

    discovery = discover_objects(token_box, refresh=args.refresh or args.refresh_discovery) if not args.render_only else (cache_load("objects_discovered.json") or [])

    accessible = [d for d in discovery if d.get("status") == 200 and d.get("name") != "_schemas_endpoint"]
    accessible_by_name = {d["name"]: d for d in accessible}

    standard_order = [n for n in OBJECT_ORDER if n in accessible_by_name]
    custom_names = [d["name"] for d in accessible if d.get("is_custom")]
    ordered = standard_order + custom_names

    objects_model: list[dict] = []
    grand_property_count = 0
    for name in ordered:
        obj_meta = accessible_by_name[name]
        props = fetch_properties(name, token_box, refresh=args.refresh) if not args.render_only else (cache_load(f"props_{name}.json") or [])
        if not props:
            continue
        live_props = [p for p in props if not p.get("archived") and not p.get("hidden")]
        live_names = [p["name"] for p in live_props]
        total = None
        pop = {}
        if not args.skip_pop and not args.render_only:
            total = fetch_total_count(name, token_box, refresh=args.refresh or args.refresh_pop)
            pop = fetch_pop_counts(name, live_names, token_box, refresh=args.refresh or args.refresh_pop)
        else:
            cached_count = cache_load(f"count_{name}.json") or {}
            total = cached_count.get("total")
            pop = cache_load(f"pop_{name}.json") or {}
        section = build_object_section(obj_meta, props, pop, total)
        objects_model.append(section)
        grand_property_count += section["property_count_visible"]
        print(f"[model] {name}: {section['property_count_visible']} live props, {len(section['sections'])} groups, total records={total}")

    now = datetime.now(timezone.utc)
    run_meta = {
        "generated_at": now.strftime("%Y-%m-%d %H:%M UTC"),
        "generated_at_date": now.strftime("%Y-%m-%d"),
        "object_count": len(objects_model),
        "property_count": grand_property_count,
    }
    model = {"objects": objects_model, "discovery": discovery}
    print(f"[render] writing {OUTPUT_DOCX}")
    render_docx(model, run_meta, OUTPUT_DOCX)
    print(f"[done] {OUTPUT_DOCX}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
